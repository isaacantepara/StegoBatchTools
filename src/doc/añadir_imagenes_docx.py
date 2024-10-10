import os
from docx import Document
from docx.shared import Inches
from PIL import Image
import tempfile

def convert_and_add_images_to_document(folder_path, output_filename):
    # Crear un nuevo documento
    doc = Document()

    # Obtener la lista de archivos de imagen en la carpeta
    image_files = [f for f in os.listdir(folder_path) if f.lower().endswith(('.jpeg', '.jpg', '.png', '.gif', '.bmp'))]

    with tempfile.TemporaryDirectory() as temp_dir:
        for image_file in image_files:
            image_path = os.path.join(folder_path, image_file)
            try:
                # Abrir la imagen con Pillow
                with Image.open(image_path) as img:
                    # Crear un nombre temporal para la imagen convertida
                    temp_image_path = os.path.join(temp_dir, f"temp_{image_file}.png")
                    
                    # Convertir y guardar la imagen como PNG
                    img.save(temp_image_path, "PNG")
                    
                    # Añadir la imagen convertida al documento
                    doc.add_picture(temp_image_path, width=Inches(6))
                    doc.add_paragraph(f"Imagen: {image_file}")
                    print(f"Imagen añadida con éxito: {image_file}")
            except Exception as e:
                print(f"Error al procesar la imagen {image_file}: {str(e)}")
                doc.add_paragraph(f"Error al procesar imagen: {image_file}")

    # Guardar el documento
    try:
        doc.save(output_filename)
        print(f"Documento creado con éxito: {output_filename}")
    except Exception as e:
        print(f"Error al guardar el documento: {str(e)}")

# Ruta de la carpeta que contiene las imágenes
folder_path = '/home/pythonesso/Downloads/mierda/peces'

# Nombre del archivo de salida
output_filename = 'peces.docx'

# Llamar a la función
convert_and_add_images_to_document(folder_path, output_filename)