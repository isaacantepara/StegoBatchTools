from docx import Document
from faker import Faker
import random
import os
from datetime import datetime
import wikipedia

# Inicializar Faker
fake = Faker()

# Lista de nombres de animales (ampliada) en inglés
animales = [
    'Lion', 'Asian Elephant', 'Giraffe', 'Tiger', 'Giant Panda', 
    'Grevy\'s Zebra', 'Red Kangaroo', 'Hippopotamus', 'White Rhinoceros', 
    'Bornean Orangutan', 'Nile Crocodile', 'American Flamingo', 'Emperor Penguin', 
    'RedFox', 'Gray Wolf', 'Brown Bear', 'Burmese Python', 
    'River Otter', 
    'Puma', 'Eagle Owl', 'Kestrel', 'Western Gorilla', 'Common Chimpanzee', 
    'Ring-tailed Lemur', 'Capuchin Monkey', 'Koala', 'Bottlenose Dolphin', 'Humpback Whale', 
    'Great White Shark', 'Loggerhead Sea Turtle', 'Boa Constrictor', 'Goanna', 
    'Green Iguana', 'Bullfrog', 'Common Toad', 'Monarch Butterfly', 'Honey Bee', 
    'Silkworm', 'Dung Beetle', 'Woodpecker', 'Bald Eagle', 
    'Peregrine Falcon', 'Herring Gull',
    'Canada Goose', 'Mute Swan', 
    'Common Quail', 'Rock Dove', 'Indian Peafowl', 'Common Pheasant', 
    'Domestic Dog', 'Domestic Cat', 'European Rabbit', 'Horse', 
    'Domestic Pig', 'Sheep', 'Goat', 'Donkey', 'Dromedary Camel', 
    'Northern Elephant Seal', 'Harbor Seal', 'West Indian Manatee', 'Beaver', 
    'Gray Squirrel', 'House Mouse', 'European Hedgehog', 'Bat', 
    'Anna\'s Hummingbird', 'African Grey Parrot', 'Blue-and-yellow Macaw', 'Domestic Canary', 
    'Crayfish', 'American Lobster', 'Common Octopus', 'Giant Squid', 
    'Jellyfish', 'Starfish', 'Goldfish', 'Hammerhead Shark', 
    'Blue Whale', 'Bottlenose Dolphin', 'Seahorse', 'Red Coral', 
    'Green Algae', 'Plankton', 'Bacteria', 'Mushroom', 
    'Houseplant', 'Oak Tree', 'Rose'
]
# Función para crear un informe para cada animal
def generar_informe(animal):
    # Crear un nuevo documento
    doc = Document()
    
    # Obtener la fecha actual
    fecha_actual = datetime.now()

    # Título del informe
    doc.add_heading(f'Informe de {animal}', 0)
    
    # Información general
    doc.add_heading('Información General', level=1)
    visitas = random.randint(5000, 20000)
    doc.add_paragraph(f"El zoológico ha recibido un total de {visitas} visitantes en la fecha del informe.")
    doc.add_paragraph(f"Este informe fue elaborado por: {fake.name()}") # Nombre del recolector
    
    # Información sobre el animal
    doc.add_heading('Estado del Animal', level=1)
    try:
        # Busca información sobre el animal en Wikipedia (en inglés)
        page = wikipedia.page(animal)
        
        # Extrae información relevante
        nombre_cientifico = page.title
        estado_conservacion = page.summary.split(" ")[0]  
        descripcion = page.content  # Obtiene todo el contenido de la página
        
        # Asigna información a la sección del documento
        doc.add_paragraph(f"Nombre científico: {nombre_cientifico}")
        doc.add_paragraph(f"Estado de conservación: {estado_conservacion}")
        doc.add_paragraph(descripcion)

    except wikipedia.exceptions.PageError:
        # Maneja el error si no se encuentra información
        doc.add_paragraph(f"No se encontró información sobre {animal} en Wikipedia.")

    # Información adicional (datos ficticios)
    doc.add_heading('Información adicional', level=1)
    doc.add_paragraph(f"Peso: {random.randint(10, 500)} kg")
    doc.add_paragraph(f"Altura: {random.randint(1, 5)} m")
    doc.add_paragraph(f"Longitud: {random.randint(1, 10)} m")
    doc.add_paragraph(f"Velocidad máxima: {random.randint(10, 100)} km/h")
    doc.add_paragraph(f"Esperanza de vida: {random.randint(5, 50)} años")
    doc.add_paragraph(f"Número de crías por camada: {random.randint(1, 10)}")
    doc.add_paragraph(f"Tiempo de gestación: {random.randint(1, 12)} meses")

    # Agregar una sección con datos adicionales
    doc.add_heading('Datos adicionales', level=1)
    doc.add_paragraph(f"Distribución geográfica: {fake.sentence(nb_words=5)}")
    doc.add_paragraph(f"Dieta: {fake.sentence(nb_words=3)}")
    doc.add_paragraph(f"Amenazas: {fake.sentence(nb_words=5)}")
    doc.add_paragraph(f"Medidas de conservación: {fake.sentence(nb_words=5)}")

    # Agregar una sección con información sobre el cuidado en el zoológico
    doc.add_heading('Cuidado en el zoológico', level=1)
    doc.add_paragraph(f"Hábitat del zoológico: {fake.sentence(nb_words=5)}")
    doc.add_paragraph(f"Dieta del zoológico: {fake.sentence(nb_words=5)}")
    doc.add_paragraph(f"Programa de enriquecimiento ambiental: {fake.sentence(nb_words=5)}")
    doc.add_paragraph(f"Programa de reproducción: {fake.sentence(nb_words=5)}")

    # Guardar el documento
    dia_random = random.randint(1, 31)  # Generar un día aleatorio
    nombre_archivo = f"../pendriver/informes/Informe_{animal.replace(' ', '_')}_{dia_random}_{fecha_actual.month}.docx"
    doc.save(nombre_archivo)
    print(f"Informe guardado como: {nombre_archivo}")

# Directorio de salida
output_dir = "../pendriver/informes"
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

# Generar informes para cada animal
for animal in animales:
    generar_informe(animal)